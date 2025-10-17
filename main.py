from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse, Response
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from io import BytesIO
from typing import List, Dict
from models import Ship, CreateShip
from bd import connect_db

app = FastAPI(title="Ship Management Web")

@app.get("/", response_class=HTMLResponse)
async def read_root():
    with open("templates/index.html", "r", encoding="utf-8") as f:
        html_content = f.read()
    return HTMLResponse(content=html_content)

@app.get("/ships")
async def get_ships() -> List:
    conn = connect_db()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT * FROM ships")
        ships_data = cursor.fetchall()
        
        ships = []
        for ship_data in ships_data:
            ship = Ship(id=ship_data[0],
                        name=ship_data[1],
                        displacement=ship_data[2],
                        port=ship_data[3],
                        captain=ship_data[4],
                        berth_num=ship_data[5],
                        target=ship_data[6])
            
            ships.append(ship)

        return ships


    except Exception as e:
        print(f"Ошибка в получении данных из базы данных, {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        cursor.close()
        conn.close()



@app.get("/ships/download_word")
async def export_ships_to_word():
    conn = connect_db()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT * FROM ships")
        ships = cursor.fetchall()
    except Exception as e:
        print(f"Ошибка получения данных их базы данных {e}")
        raise HTTPException(status_code=500, detail="Database error")

    try:
        if not ships:
            raise HTTPException(status_code=404, detail="No ships available for export")
            
        doc = Document()
        doc.add_heading('Ship Fleet Report', 0)
        doc.add_paragraph(f"Total ships in fleet: {len(ships)}")
        doc.add_paragraph()

        for i, ship in enumerate(ships, 1):
            if i > 1:
                doc.add_paragraph("-" * 50)
                doc.add_paragraph()
            
            doc.add_heading(f'Ship #{ship[0]}: {ship[1]}', level=2)
            doc.add_paragraph(f"ID: {ship[0]}")
            doc.add_paragraph(f"Ship Name: {ship[1]}")
            doc.add_paragraph(f"Displacement: {ship[2]} tons")
            doc.add_paragraph(f"Home Port: {ship[3]}")
            doc.add_paragraph(f"Captain: {ship[4]}")
            doc.add_paragraph(f"Berth Number: {ship[5]}")
            doc.add_paragraph(f"Destination: {ship[6]}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': 'attachment; filename="ship_fleet_report.docx"'}
        )
    except Exception as e:
        print(f"Error in export_ships_to_word: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")
    finally:
        cursor.close()
        conn.close()
    
@app.get("/ships/download_excel")
async def export_ships_to_excel():
    conn = connect_db()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT * FROM ships")
        ships = cursor.fetchall()
    except Exception as e:
        print(f"Ошибка получения данных их базы данных {e}")
        raise HTTPException(status_code=500, detail="Database error")

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Ship Information"

        # Добавляем название таблицы
        ws['A1'] = f"Ship Report"
        title_cell = ws['A1']
        title_cell.font = Font(size=14, bold=True)
        ws.merge_cells('A1:G1')  # Объединяем ячейки для заголовка
        ws.row_dimensions[1].height = 25  # Высота строки для заголовка

        current_row = 3

        # Добавляем заголовки таблицы
        headers = ["ID", "Ship Name", "Displacement", "Home Port", "Captain", "Berth Number", "Destination"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

        for ship in ships:
            # Добавляем данные корабля
            current_row += 1
            ship_data = [
                ship[0],
                ship[1],
                f"{ship[2]} tons",
                ship[3],
                ship[4],
                ship[5],
                ship[6]
            ]

            for col, value in enumerate(ship_data, 1):
                cell = ws.cell(row=current_row, column=col, value=value)
                cell.border = Border(
                    left=Side(style='thin'),
                    right=Side(style='thin'),
                    top=Side(style='thin'),
                    bottom=Side(style='thin')
                )

            # Выравнивание заголовка по центру
            ws['A1'].alignment = Alignment(horizontal='center', vertical='center')

            # Выравнивание всех ячеек по центру
            for row in ws.iter_rows(min_row=3, max_row=current_row):
                for cell in row:
                    cell.alignment = Alignment(horizontal='center', vertical='center')

        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        filename = f"ships_report.xlsx"

        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        print(f"Error in export_ship_to_excel: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating Excel report: {str(e)}")
    finally:
        cursor.close()
        conn.close()

@app.get("/ships/{ship_id}")
async def get_ship(ship_id: int) -> Ship:
    conn = connect_db()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT * FROM ships WHERE id = %s", (ship_id,))
        ship_data = cursor.fetchone()
        if not ship:
            raise HTTPException(status_code=404, detail="Ship not found")

        ship = Ship(
            id=ship_data[0],
            name=ship_data[1],
            displacement=ship_data[2],
            port=ship_data[3],
            captain=ship_data[4],
            berth_num=ship_data[5],
            target=ship_data[6]
        )

        return ship

    except Exception as e:
        print(f"Ошибка получения данных их базы данных {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        cursor.close()
        conn.close()

@app.post("/ships")
async def add_ship(ship: CreateShip) -> Ship:
    conn = connect_db()
    cursor = conn.cursor()

    try:
        cursor.execute("""
            INSERT INTO ships (name, displacement, port, captain, berth_num, target)
            VALUES (%s, %s, %s, %s, %s, %s)
            RETURNING id, name, displacement, port, captain, berth_num, target
        """, (ship.name, 
              ship.displacement, 
              ship.port, ship.captain,
              ship.berth_num,
              ship.target))
        
        new_ship_data = cursor.fetchone()
        conn.commit()

        ship_response = Ship(
            id=new_ship_data[0],
            name=new_ship_data[1],
            displacement=new_ship_data[2],
            port=new_ship_data[3],
            captain=new_ship_data[4],
            berth_num=new_ship_data[5],
            target=new_ship_data[6]
        )
        
        return ship_response
    
    except Exception as e:
        conn.rollback()
        print(f"Ошибка подключения к базе данных {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        cursor.close()
        conn.close()
    

@app.put("/ships/{ship_id}")
async def edit_ship(ship_id: int, ship_data: CreateShip) -> Ship:
    conn = connect_db()
    if conn is None:
        raise HTTPException(status_code=500, detail="Database connection failed")
    
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT * FROM ships WHERE id = %s", (ship_id,))
        if not cursor.fetchone():
            raise HTTPException(status_code=404, detail="Ship not found")
        
        cursor.execute("""
            UPDATE ships 
            SET name = %s,
                displacement = %s,
                port = %s,
                captain = %s,
                berth_num = %s,
                target = %s
            WHERE id = %s
            RETURNING id, name, displacement, port, captain, berth_num, target
        """, (
            ship_data.name,
            ship_data.displacement,
            ship_data.port,
            ship_data.captain,
            ship_data.berth_num,
            ship_data.target,
            ship_id
        ))
        
        update_ship = cursor.fetchone()
        conn.commit()

        edited_ship = Ship(
            id=update_ship[0],
            name=update_ship[1],
            displacement=update_ship[2],
            port=update_ship[3],
            captain=update_ship[4],
            berth_num=update_ship[5],
            target=update_ship[6]
        )
        
        return edited_ship
    
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        print(f"Ошибка при обновлении корабля: {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        cursor.close()
        conn.close()

@app.delete("/ships/{ship_id}")  
async def delete_ship(ship_id: int) -> Dict:
    conn = connect_db()
    cursor = conn.cursor()

    try:
        cursor.execute("SELECT * FROM ships WHERE id = %s", (ship_id,))
        if not cursor.fetchone():
            raise HTTPException(status_code=404, detail="Ship not found")
        
        cursor.execute("DELETE FROM ships WHERE id = %s", (ship_id,))
        conn.commit()

        return {"data": "Ship deleted"}
    
    except Exception as e:
        conn.rollback()
        print(f"Ошибка в удалении данных из базы данных, {e}")
        raise HTTPException(status_code=500, detail="Database error")
    finally:
        cursor.close()
        conn.close()

@app.get("/ships/download_word/{ship_id}")
async def export_ship_to_word(ship_id: int):
    conn = connect_db()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT * FROM ships WHERE id = %s", (ship_id,))
        ship = cursor.fetchone()
        if not ship:
            raise HTTPException(status_code=404, detail="Ship not found")

        doc = Document()
        doc.add_heading(f'Ship Report: {ship[1]}', 0)
        
        # Данные корабля
        ship_data = [
            ("ID", str(ship[0])),
            ("Ship Name", ship[1]),
            ("Displacement", f"{ship[2]} tons"),
            ("Home Port", ship[3]),
            ("Captain", ship[4]),
            ("Berth Number", str(ship[5])),
            ("Destination", ship[6])
        ]
        
        for field, value in ship_data:
            p = doc.add_paragraph()
            p.add_run(f"{field}: ").bold = True
            p.add_run(value)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"ship_{ship[0]}_report.docx"
        
        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )
        
    except Exception as e:
        print(f"Error in export_ship_to_word: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating ship report: {str(e)}")
    finally:
        cursor.close()
        conn.close()

@app.get("/ships/download_excel/{ship_id}")
async def export_ship_to_excel(ship_id: int):
    conn = connect_db()
    cursor = conn.cursor()
    try:
        cursor.execute("SELECT * FROM ships WHERE id = %s", (ship_id,))
        ship = cursor.fetchone()
        if not ship:
            raise HTTPException(status_code=404, detail="Ship not found")
        
        wb = Workbook()
        ws = wb.active
        ws.title = "Ship Information"
        # Добавляем название таблицы
        ws['A1'] = f"Ship Report: {ship[1]}"
        title_cell = ws['A1']
        title_cell.font = Font(size=14, bold=True)
        ws.merge_cells('A1:G1')  # Объединяем ячейки для заголовка
        ws.row_dimensions[1].height = 25  # Высота строки для заголовка
        current_row = 3
        # Добавляем заголовки таблицы
        headers = ["ID", "Ship Name", "Displacement", "Home Port", "Captain", "Berth Number", "Destination"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=current_row, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
        # Добавляем данные корабля
        current_row += 1
        ship_data = [
            ship[0],
            ship[1],
            f"{ship[2]} tons",
            ship[3],
            ship[4],
            ship[5],
            ship[6]
        ]
        for col, value in enumerate(ship_data, 1):
            cell = ws.cell(row=current_row, column=col, value=value)
            cell.border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
        # Выравнивание заголовка по центру
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        # Выравнивание всех ячеек по центру
        for row in ws.iter_rows(min_row=3, max_row=current_row):
            for cell in row:
                cell.alignment = Alignment(horizontal='center', vertical='center')
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        filename = f"ship_{ship[0]}_report.xlsx"
        return Response(
            content=buffer.getvalue(),
            media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            headers={'Content-Disposition': f'attachment; filename="{filename}"'}
        )

    except Exception as e:
        print(f"Error in export_ship_to_excel: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating Excel report: {str(e)}")
    finally:
        cursor.close()
        conn.close()